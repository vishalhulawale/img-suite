"""DOCX converter — pixel-perfect PDF to Word conversion.

Two layout modes:
  - FLOW: Normal paragraphs, tables, images for simple layouts.
          Preserves fonts, sizes, colors, spacing, alignment, lists, headings.
  - FIXED: Absolute positioning via text boxes for complex layouts.
           Every element placed at its exact PDF coordinate.
  - AUTO: Per-page decision based on layout complexity.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

from app.pdf_converter.config import DocxLayoutMode, PipelineConfig
from app.pdf_converter.models import (
    BorderStyle, Color, Document, DrawingElement, ElementType, FontInfo,
    ImageElement, ListType, Page, PageElement, PageType, Rect, TableCell, TableElement,
    TextAlignment, TextBlock, TextLine, TextSpan,
)

logger = logging.getLogger(__name__)

# Points to EMU conversion (1 point = 12700 EMU)
PT_TO_EMU = 12700
# Inches to EMU
IN_TO_EMU = 914400


def convert_to_docx(doc: Document, output_path: str, config: PipelineConfig) -> str:
    """Convert internal document model to DOCX file.

    Returns path to output file.
    """
    docx = DocxDocument()

    # Set default styles
    _setup_default_styles(docx, doc, config)

    last_page_idx = len(doc.pages) - 1
    for page_idx, page in enumerate(doc.pages):
        # Decide layout mode first so we can configure section correctly
        mode = _decide_layout_mode(page, config)
        is_last = page_idx == last_page_idx

        # Add section break for each page (except first)
        if page_idx > 0:
            _add_section_break(docx, page, config, mode)
        else:
            _configure_section(docx.sections[0], page, config, mode)

        if mode == DocxLayoutMode.FIXED:
            _render_page_fixed(docx, page, doc, config)
        else:
            _render_page_flow(docx, page, doc, config, is_last_page=is_last)

    docx.save(output_path)
    logger.info("DOCX saved to: %s", output_path)
    return output_path


# ---------------------------------------------------------------------------
# Setup
# ---------------------------------------------------------------------------

def _setup_default_styles(docx: DocxDocument, doc: Document, config: PipelineConfig):
    """Configure default document styles."""
    style = docx.styles["Normal"]
    font = style.font
    font.size = Pt(11)
    font.name = "Calibri"

    # Eliminate default paragraph spacing to match PDF layout
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    # Set narrow margins initially (will be overridden per section)
    section = docx.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def _configure_section(section, page: Page, config: PipelineConfig, mode: DocxLayoutMode = DocxLayoutMode.FLOW):
    """Configure a DOCX section to match PDF page dimensions."""
    section.page_width = Emu(int(page.width * PT_TO_EMU))
    section.page_height = Emu(int(page.height * PT_TO_EMU))

    has_background = _page_has_background(page)

    if mode == DocxLayoutMode.FIXED or has_background:
        section.top_margin = Emu(0)
        section.bottom_margin = Emu(0)
        section.left_margin = Emu(0)
        section.right_margin = Emu(0)
    else:
        section.top_margin = Emu(int(page.margin_top * PT_TO_EMU))
        section.bottom_margin = Emu(int(page.margin_bottom * PT_TO_EMU))
        section.left_margin = Emu(int(page.margin_left * PT_TO_EMU))
        section.right_margin = Emu(int(page.margin_right * PT_TO_EMU))

    # Orientation
    if page.width > page.height:
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.orientation = WD_ORIENT.PORTRAIT

    # Columns
    if page.num_columns > 1:
        _set_columns(section, page.num_columns, page.column_gap)


def _set_columns(section, num_cols: int, gap_pts: float):
    """Set number of columns on a section."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), str(int(gap_pts * 20)))  # twips = pts * 20
    cols.set(qn("w:equalWidth"), "1")


def _add_section_break(docx: DocxDocument, page: Page, config: PipelineConfig, mode: DocxLayoutMode = DocxLayoutMode.FLOW):
    """Add a new section with page break."""
    new_section = docx.add_section()
    _configure_section(new_section, page, config, mode)


def _decide_layout_mode(page: Page, config: PipelineConfig) -> DocxLayoutMode:
    """Decide whether to use flow or fixed layout for a page."""
    if config.docx.layout_mode != DocxLayoutMode.AUTO:
        return config.docx.layout_mode

    # Slide-like pages (landscape with background imagery) need absolute positioning
    if page.width > page.height and _page_has_background(page):
        return DocxLayoutMode.FIXED

    # Heuristic: if page has overlapping elements or >2 columns, use FIXED
    if page.num_columns > 2:
        return DocxLayoutMode.FIXED

    # Check for overlapping text blocks
    text_bboxes = [e.bbox for e in page.elements if e.element_type == ElementType.TEXT_BLOCK]
    for i, b1 in enumerate(text_bboxes):
        for b2 in text_bboxes[i + 1:]:
            if b1.intersects(b2):
                return DocxLayoutMode.FIXED

    # Check if elements are in non-standard positions (sidebars, etc.)
    for region in page.layout_regions:
        from app.pdf_converter.models import LayoutRegionType
        if region.region_type == LayoutRegionType.SIDEBAR:
            return DocxLayoutMode.FIXED

    return DocxLayoutMode.FLOW


# ---------------------------------------------------------------------------
# Flow layout rendering
# ---------------------------------------------------------------------------

def _render_page_flow(
    docx: DocxDocument,
    page: Page,
    doc: Document,
    config: PipelineConfig,
    is_last_page: bool = False,
):
    """Render page elements as normal document flow (paragraphs, tables, images)."""
    _render_page_backgrounds(docx, page, config)

    has_background = _page_has_background(page)
    if has_background:
        margin_top = margin_bottom = margin_left = margin_right = 0.0
    else:
        margin_top = page.margin_top
        margin_bottom = page.margin_bottom
        margin_left = page.margin_left
        margin_right = page.margin_right

    SECTION_BREAK_OVERHEAD = 22.0
    content_width_pts = page.width - margin_left - margin_right
    content_height_pts = page.height - margin_top - margin_bottom - SECTION_BREAK_OVERHEAD

    # Sort elements by reading order
    elements = sorted(page.elements, key=lambda e: (e.reading_order, e.bbox.y0, e.bbox.x0))

    text_height = 0.0
    image_height = 0.0
    for elem in elements:
        if elem.is_background:
            continue
        if elem.element_type == ElementType.TEXT_BLOCK:
            max_font = max(
                (s.font.size for l in elem.element.lines for s in l.spans),
                default=12.0,
            )
            n_lines = sum(len(b.lines) for b in [elem.element])
            text_height += max(elem.bbox.height, max_font * 1.15 * n_lines)
        elif elem.element_type == ElementType.IMAGE:
            image_height += elem.bbox.height

    total_height = text_height + image_height
    image_scale = 1.0
    if total_height > content_height_pts and image_height > 0:
        available_for_images = max(content_height_pts - text_height, content_height_pts * 0.25)
        image_scale = min(1.0, available_for_images / image_height)

    for elem in elements:
        if elem.is_background:
            continue

        if elem.element_type == ElementType.TEXT_BLOCK:
            _add_text_block_flow(docx, elem.element, config)
        elif elem.element_type == ElementType.TABLE:
            _add_table_flow(docx, elem.element, config)
        elif elem.element_type == ElementType.IMAGE:
            _add_image_flow(docx, elem.element, config, content_width_pts, image_scale)


def _add_text_block_flow(
    docx: DocxDocument,
    block: TextBlock,
    config: PipelineConfig,
):
    """Add a text block as a DOCX paragraph."""
    style_name = None
    if block.heading_level > 0 and config.docx.map_heading_styles:
        style_name = f"Heading {min(block.heading_level, 9)}"

    if block.list_type == ListType.BULLET:
        pass
    elif block.list_type == ListType.NUMBERED:
        style_name = "List Number"

    try:
        para = docx.add_paragraph(style=style_name)
    except KeyError:
        para = docx.add_paragraph()

    # Alignment
    align_map = {
        TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para.alignment = align_map.get(block.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Paragraph formatting
    pf = para.paragraph_format
    pf.space_before = Pt(block.space_before)
    pf.space_after = Pt(block.space_after)

    if block.indent_first > 0:
        pf.first_line_indent = Emu(int(block.indent_first * PT_TO_EMU))
    if block.indent_left > 0:
        pf.left_indent = Emu(int(block.indent_left * PT_TO_EMU))
    if block.indent_right > 0:
        pf.right_indent = Emu(int(block.indent_right * PT_TO_EMU))

    # Line spacing
    if block.line_spacing > 0:
        if 0.9 <= block.line_spacing <= 1.1:
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif 1.4 <= block.line_spacing <= 1.6:
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif 1.9 <= block.line_spacing <= 2.1:
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = block.line_spacing

    # Background color
    if block.background_color and config.docx.preserve_background_colors:
        _set_paragraph_shading(para, block.background_color)

    # Add runs for each span
    for line_idx, line in enumerate(block.lines):
        if line_idx > 0 and config.docx.preserve_line_breaks:
            if para.runs:
                para.runs[-1].add_break()

        for span in line.spans:
            _add_span_as_run(para, span, config)


def _add_span_as_run(para, span: TextSpan, config: PipelineConfig):
    """Add a TextSpan as a DOCX Run with full formatting."""
    run = para.add_run(span.text)
    font = run.font

    # Font name
    if config.docx.preserve_exact_fonts:
        font.name = span.font.name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), span.font.name)
        rFonts.set(qn("w:hAnsi"), span.font.name)
        rFonts.set(qn("w:cs"), span.font.name)

    # Size
    font.size = Pt(span.font.size)

    # Bold / Italic
    font.bold = span.font.bold
    font.italic = span.font.italic

    # Underline / Strikethrough
    font.underline = span.font.underline
    font.strike = span.font.strikethrough

    # Color
    color = span.font.color
    if color:
        font.color.rgb = RGBColor(*color.to_rgb_int())

    # Superscript / Subscript
    font.superscript = span.font.superscript
    font.subscript = span.font.subscript

    # Highlight
    if span.font.highlight_color:
        _set_run_highlight(run, span.font.highlight_color)

    # Character spacing
    if span.char_spacing and abs(span.char_spacing) > 0.1:
        _set_char_spacing(run, span.char_spacing)


def _set_paragraph_shading(para, color: Color):
    """Set paragraph background/shading color."""
    hex_color = color.to_hex().lstrip("#")
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    para._element.get_or_add_pPr().append(shading_elm)


def _set_run_highlight(run, color: Color):
    """Set run highlight color (approximate to available DOCX highlight colors)."""
    r, g, b = color.to_rgb_int()
    if r > 200 and g > 200 and b < 100:
        highlight = "yellow"
    elif r < 100 and g > 200 and b < 100:
        highlight = "green"
    elif r < 100 and g < 100 and b > 200:
        highlight = "blue"
    elif r > 200 and g < 100 and b < 100:
        highlight = "red"
    elif r > 200 and g > 100 and b < 100:
        highlight = "darkYellow"
    else:
        highlight = "yellow"

    rPr = run._element.get_or_add_rPr()
    highlight_elm = OxmlElement("w:highlight")
    highlight_elm.set(qn("w:val"), highlight)
    rPr.append(highlight_elm)


def _set_char_spacing(run, spacing_pts: float):
    """Set character spacing on a run."""
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(int(spacing_pts * 20)))
    rPr.append(spacing)


# ---------------------------------------------------------------------------
# Table rendering (flow mode)
# ---------------------------------------------------------------------------

def _add_table_flow(
    docx: DocxDocument,
    table: TableElement,
    config: PipelineConfig,
):
    """Add a table to the document."""
    if table.num_rows == 0 or table.num_cols == 0:
        return

    docx_table = docx.add_table(rows=table.num_rows, cols=table.num_cols)
    docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    docx_table.autofit = False

    # Set column widths
    for col_idx, width_pts in enumerate(table.col_widths):
        if col_idx < len(docx_table.columns):
            docx_table.columns[col_idx].width = Emu(int(width_pts * PT_TO_EMU))

    # Set row heights
    for row_idx, height_pts in enumerate(table.row_heights):
        if row_idx < len(docx_table.rows):
            tr = docx_table.rows[row_idx]._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(int(height_pts * 20)))  # twips
            trHeight.set(qn("w:hRule"), "exact")
            trPr.append(trHeight)

    # Track which cells are covered by spans
    covered: set[tuple[int, int]] = set()

    for cell in table.cells:
        if (cell.row_index, cell.col_index) in covered:
            continue

        if cell.row_index >= table.num_rows or cell.col_index >= table.num_cols:
            continue

        docx_cell = docx_table.cell(cell.row_index, cell.col_index)

        # Handle cell spanning
        if cell.row_span > 1 or cell.col_span > 1:
            end_row = min(cell.row_index + cell.row_span - 1, table.num_rows - 1)
            end_col = min(cell.col_index + cell.col_span - 1, table.num_cols - 1)

            try:
                merge_cell = docx_table.cell(end_row, end_col)
                docx_cell = docx_cell.merge(merge_cell)
            except Exception:
                pass

            for r in range(cell.row_index, end_row + 1):
                for c in range(cell.col_index, end_col + 1):
                    if (r, c) != (cell.row_index, cell.col_index):
                        covered.add((r, c))

        # Cell content
        _populate_cell(docx_cell, cell, config)

        # Cell background
        if cell.background_color:
            _set_cell_shading(docx_cell, cell.background_color)

        # Cell borders
        _set_cell_borders(docx_cell, cell)

        # Vertical alignment
        va_map = {
            "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
            "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        }
        docx_cell.vertical_alignment = va_map.get(
            cell.vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.TOP
        )

    # Header row repeat
    if table.has_header_row:
        _set_header_row(docx_table)


def _populate_cell(docx_cell, cell: TableCell, config: PipelineConfig):
    """Fill a DOCX cell with content from the model."""
    if docx_cell.paragraphs:
        docx_cell.paragraphs[0].clear()

    for block_idx, block in enumerate(cell.content):
        if block_idx == 0 and docx_cell.paragraphs:
            para = docx_cell.paragraphs[0]
        else:
            para = docx_cell.add_paragraph()

        for line_idx, line in enumerate(block.lines):
            if line_idx > 0:
                if para.runs:
                    para.runs[-1].add_break()
            for span in line.spans:
                _add_span_as_run(para, span, config)


def _set_cell_shading(docx_cell, color: Color):
    """Set cell background color."""
    hex_color = color.to_hex().lstrip("#")
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    docx_cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(docx_cell, cell: TableCell):
    """Set cell border styles."""
    tc = docx_cell._tc
    tcPr = tc.get_or_add_tcPr()

    borders = OxmlElement("w:tcBorders")

    for side, border in [
        ("top", cell.border_top),
        ("bottom", cell.border_bottom),
        ("left", cell.border_left),
        ("right", cell.border_right),
    ]:
        if border.style == BorderStyle.NONE:
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:val"), "none")
            borders.append(elem)
        else:
            style_map = {
                BorderStyle.SOLID: "single",
                BorderStyle.DASHED: "dashed",
                BorderStyle.DOTTED: "dotted",
                BorderStyle.DOUBLE: "double",
            }
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:val"), style_map.get(border.style, "single"))
            elem.set(qn("w:sz"), str(int(border.width * 8)))
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), border.color.to_hex().lstrip("#"))
            borders.append(elem)

    tcPr.append(borders)


def _set_header_row(docx_table):
    """Mark first row as header (repeats across pages)."""
    try:
        tr = docx_table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Image rendering (flow mode)
# ---------------------------------------------------------------------------

def _add_image_flow(
    docx: DocxDocument,
    image: ImageElement,
    config: PipelineConfig,
    content_width_pts: float = 468.0,
    height_scale: float = 1.0,
):
    """Add an image as an inline element."""
    try:
        img_stream = io.BytesIO(image.image_data)

        display_width_in = image.bbox.width / 72.0
        display_height_in = image.bbox.height / 72.0

        if height_scale < 1.0:
            display_width_in *= height_scale
            display_height_in *= height_scale

        max_width_in = content_width_pts / 72.0
        if display_width_in > max_width_in:
            scale = max_width_in / display_width_in
            display_width_in *= scale
            display_height_in *= scale

        para = docx.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(img_stream, width=Inches(display_width_in))

    except Exception as e:
        logger.warning("Failed to add image: %s", e)


# ---------------------------------------------------------------------------
# Fixed layout rendering (absolute positioning)
# ---------------------------------------------------------------------------

def _render_page_fixed(
    docx: DocxDocument,
    page: Page,
    doc: Document,
    config: PipelineConfig,
):
    """Render page with absolute positioning using text boxes."""
    _render_page_backgrounds(docx, page, config)

    elements = sorted(page.elements, key=lambda e: (e.z_order, e.reading_order))

    image_bboxes = [
        e.bbox for e in elements
        if e.element_type == ElementType.IMAGE and not e.is_background
    ]

    bg_draw_elems = [
        e for e in elements
        if e.element_type == ElementType.DRAWING and e.is_background
        and isinstance(e.element, DrawingElement) and e.element.fill_color is not None
    ]

    has_content = False
    for elem in elements:
        if elem.is_background:
            continue

        if _is_occluded_by_bg_drawing(elem, bg_draw_elems):
            continue

        if elem.element_type == ElementType.TEXT_BLOCK:
            _add_text_block_fixed(docx, elem, page, config)
            has_content = True
        elif elem.element_type == ElementType.IMAGE:
            _add_image_fixed(docx, elem, page, config)
            has_content = True
        elif elem.element_type == ElementType.TABLE:
            _add_table_flow(docx, elem.element, config)
            has_content = True
        elif elem.element_type == ElementType.DRAWING:
            if _is_artifact_drawing(elem, image_bboxes):
                continue
            _add_drawing_fixed(docx, elem, page, config)
            has_content = True

    if not has_content:
        _add_zero_height_para(docx)


def _add_zero_height_para(docx: DocxDocument):
    """Add a paragraph that takes up no vertical space in Word's page flow."""
    para = docx.add_paragraph()
    _collapse_paragraph_height(para)


def _add_text_block_fixed(
    docx: DocxDocument,
    elem: PageElement,
    page: Page,
    config: PipelineConfig,
):
    """Add a text block as an absolutely positioned text box."""
    block = elem.element
    if not isinstance(block, TextBlock):
        return

    para = docx.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    pPr = para._element.get_or_add_pPr()
    framePr = OxmlElement("w:framePr")

    x_twips = int(elem.bbox.x0 * 20)
    y_twips = int(elem.bbox.y0 * 20)
    raw_width = elem.bbox.width
    max_width = max(page.width - elem.bbox.x0, raw_width)
    max_font_size = max(
        (s.font.size for l in block.lines for s in l.spans),
        default=12.0,
    )
    width_factor = 1.25 if max_font_size >= 18 else 1.15
    w_twips = int(min(raw_width * width_factor, max_width) * 20)
    h_twips = int(elem.bbox.height * 20)

    framePr.set(qn("w:w"), str(w_twips))
    framePr.set(qn("w:h"), str(h_twips))
    framePr.set(qn("w:x"), str(x_twips))
    framePr.set(qn("w:y"), str(y_twips))
    framePr.set(qn("w:hAnchor"), "page")
    framePr.set(qn("w:vAnchor"), "page")
    framePr.set(qn("w:wrap"), "none")

    pPr.append(framePr)

    for line_idx, line in enumerate(block.lines):
        if line_idx > 0 and para.runs:
            para.runs[-1].add_break()
        for span in line.spans:
            _add_span_as_run(para, span, config)


def _add_image_fixed(
    docx: DocxDocument,
    elem: PageElement,
    page: Page,
    config: PipelineConfig,
    behind_text: bool = False,
):
    """Add an image with absolute positioning."""
    image = elem.element
    if not isinstance(image, ImageElement):
        return

    try:
        img_stream = io.BytesIO(image.image_data)
        display_width_in = elem.bbox.width / 72.0
        display_height_in = elem.bbox.height / 72.0

        para = docx.add_paragraph()
        _collapse_paragraph_height(para)
        run = para.add_run()
        pic = run.add_picture(img_stream, width=Inches(display_width_in))

        inline = run._element.find(qn("w:drawing")).find(qn("wp:inline"))
        if inline is not None:
            _convert_inline_to_anchor(
                run, inline,
                x_emu=int(elem.bbox.x0 * PT_TO_EMU),
                y_emu=int(elem.bbox.y0 * PT_TO_EMU),
                cx_emu=int(elem.bbox.width * PT_TO_EMU),
                cy_emu=int(elem.bbox.height * PT_TO_EMU),
                behind_text=behind_text,
            )

    except Exception as e:
        logger.warning("Failed to add fixed image: %s", e)


def _convert_inline_to_anchor(run, inline, x_emu, y_emu, cx_emu, cy_emu, behind_text: bool = False):
    """Convert an inline drawing to an anchored (absolute position) drawing."""
    drawing = run._element.find(qn("w:drawing"))
    if drawing is None:
        return

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "0")
    anchor.set("behindDoc", "1" if behind_text else "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simplePos = OxmlElement("wp:simplePos")
    simplePos.set("x", "0")
    simplePos.set("y", "0")
    anchor.append(simplePos)

    posH = OxmlElement("wp:positionH")
    posH.set("relativeFrom", "page")
    posOffset = OxmlElement("wp:posOffset")
    posOffset.text = str(x_emu)
    posH.append(posOffset)
    anchor.append(posH)

    posV = OxmlElement("wp:positionV")
    posV.set("relativeFrom", "page")
    posOffset = OxmlElement("wp:posOffset")
    posOffset.text = str(y_emu)
    posV.append(posOffset)
    anchor.append(posV)

    extent = OxmlElement("wp:extent")
    extent.set("cx", str(cx_emu))
    extent.set("cy", str(cy_emu))
    anchor.append(extent)

    wrapNone = OxmlElement("wp:wrapNone")
    anchor.append(wrapNone)

    graphic = inline.find(qn("a:graphic"))
    if graphic is not None:
        anchor.append(graphic)

    docPr = inline.find(qn("wp:docPr"))
    if docPr is not None:
        anchor.append(docPr)

    cNvPr = inline.find(qn("wp:cNvGraphicFramePr"))
    if cNvPr is not None:
        anchor.append(cNvPr)

    drawing.remove(inline)
    drawing.append(anchor)


def _page_has_background(page: Page) -> bool:
    return (
        page.background_image is not None
        or _has_visible_background_color(page.background_color)
        or any(elem.is_background for elem in page.elements)
    )


def _has_visible_background_color(color: Color | None) -> bool:
    if color is None:
        return False
    r, g, b = color.to_rgb_int()
    return (r, g, b) != (255, 255, 255)


def _render_page_backgrounds(
    docx: DocxDocument,
    page: Page,
    config: PipelineConfig,
):
    if not config.docx.preserve_background_colors:
        return

    full_page_drawn = False
    if _has_visible_background_color(page.background_color):
        _add_background_rectangle(
            docx,
            Rect(0, 0, page.width, page.height),
            page.background_color,
        )
        full_page_drawn = True

    elements = sorted(page.elements, key=lambda e: (e.z_order, e.reading_order))

    for elem in elements:
        if elem.element_type == ElementType.IMAGE and elem.is_background:
            _add_image_fixed(docx, elem, page, config, behind_text=True)

    for elem in elements:
        if not elem.is_background or elem.element_type != ElementType.DRAWING:
            continue

        drawing = elem.element
        if not isinstance(drawing, DrawingElement) or drawing.fill_color is None:
            continue

        if full_page_drawn and _covers_most_of_page(elem.bbox, page):
            continue

        _add_background_rectangle(
            docx,
            elem.bbox,
            drawing.fill_color,
            opacity=drawing.opacity,
        )


def _add_background_rectangle(
    docx: DocxDocument,
    bbox: Rect,
    color: Color,
    opacity: float = 1.0,
):
    """Add a colored background rectangle using DrawingML anchor shape."""
    para = docx.add_paragraph()
    _collapse_paragraph_height(para)

    x_emu = int(bbox.x0 * PT_TO_EMU)
    y_emu = int(bbox.y0 * PT_TO_EMU)
    cx_emu = max(int(bbox.width * PT_TO_EMU), 1)
    cy_emu = max(int(bbox.height * PT_TO_EMU), 1)
    color_hex = color.to_hex().lstrip("#")
    shape_id = abs(hash((bbox.x0, bbox.y0, bbox.x1, bbox.y1))) % 65534 + 1

    alpha_xml = ""
    if opacity < 0.999:
        alpha_val = int(opacity * 100000)
        alpha_xml = f'<a:alpha val="{alpha_val}"/>'

    drawing_xml = (
        '<w:drawing'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        ' simplePos="0" relativeHeight="251658240" behindDoc="1"'
        ' locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx_emu}" cy="{cy_emu}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{shape_id}" name="bg_{shape_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>'
        '<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx_emu}" cy="{cy_emu}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:solidFill><a:srgbClr val="{color_hex}">{alpha_xml}</a:srgbClr></a:solidFill>'
        '<a:ln><a:noFill/></a:ln>'
        '</wps:spPr>'
        '<wps:bodyPr/>'
        '</wps:wsp>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    )
    para.add_run()._r.append(parse_xml(drawing_xml))


def _collapse_paragraph_height(para):
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)

    spacing.set(qn("w:line"), "20")
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")


def _covers_most_of_page(bbox: Rect, page: Page, threshold: float = 0.85) -> bool:
    page_area = page.width * page.height
    if page_area <= 0 or bbox.area <= 0:
        return False
    return bbox.area / page_area >= threshold


def _add_drawing_fixed(
    docx: DocxDocument,
    elem: PageElement,
    page: Page,
    config: PipelineConfig,
):
    """Add a non-background drawing element as a DrawingML anchored shape."""
    drawing = elem.element
    if not isinstance(drawing, DrawingElement):
        return
    if not drawing.fill_color and not drawing.stroke_color:
        return

    para = docx.add_paragraph()
    _collapse_paragraph_height(para)

    x_emu = int(elem.bbox.x0 * PT_TO_EMU)
    y_emu = int(elem.bbox.y0 * PT_TO_EMU)
    cx_emu = max(int(elem.bbox.width * PT_TO_EMU), 1)
    cy_emu = max(int(elem.bbox.height * PT_TO_EMU), 1)
    shape_id = abs(hash((elem.bbox.x0, elem.bbox.y0, elem.bbox.x1, elem.bbox.y1, "fg"))) % 65534 + 1

    if drawing.fill_color:
        color_hex = drawing.fill_color.to_hex().lstrip("#")
        alpha_xml = ""
        if drawing.opacity < 0.999:
            alpha_xml = f'<a:alpha val="{int(drawing.opacity * 100000)}"/>'
        fill_xml = f'<a:solidFill><a:srgbClr val="{color_hex}">{alpha_xml}</a:srgbClr></a:solidFill>'
    else:
        fill_xml = '<a:noFill/>'

    if drawing.stroke_color:
        stroke_hex = drawing.stroke_color.to_hex().lstrip("#")
        stroke_emu = max(int(drawing.stroke_width * PT_TO_EMU), 1)
        line_xml = (
            f'<a:ln w="{stroke_emu}">'
            f'<a:solidFill><a:srgbClr val="{stroke_hex}"/></a:solidFill>'
            '</a:ln>'
        )
    else:
        line_xml = '<a:ln><a:noFill/></a:ln>'

    drawing_xml = (
        '<w:drawing'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        ' simplePos="0" relativeHeight="268435456" behindDoc="0"'
        ' locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx_emu}" cy="{cy_emu}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{shape_id}" name="shape_{shape_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>'
        '<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx_emu}" cy="{cy_emu}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'{fill_xml}'
        f'{line_xml}'
        '</wps:spPr>'
        '<wps:bodyPr/>'
        '</wps:wsp>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    )
    para.add_run()._r.append(parse_xml(drawing_xml))


def _is_artifact_drawing(elem: PageElement, image_bboxes: list[Rect]) -> bool:
    """Check if a drawing is a small UI artifact fully contained within an image."""
    if elem.bbox.area <= 0:
        return True
    if elem.bbox.width > 200 or elem.bbox.height > 100:
        return False
    for img_bbox in image_bboxes:
        if img_bbox.contains(elem.bbox):
            return True
    return False


def _is_occluded_by_bg_drawing(elem: PageElement, bg_draw_elems: list[PageElement]) -> bool:
    """Check if a FG element is hidden behind an opaque background drawing."""
    if not bg_draw_elems:
        return False
    eb = elem.bbox
    if eb.area <= 0:
        return False

    for bg_elem in bg_draw_elems:
        draw = bg_elem.element
        fc = draw.fill_color
        if fc is None:
            continue
        if fc.r > 0.9 and fc.g > 0.9 and fc.b > 0.9:
            continue

        db = bg_elem.bbox
        inter = eb.intersection(db)
        if inter is None:
            continue
        overlap = inter.area / eb.area if eb.area > 0 else 0
        if overlap < 0.7:
            continue

        if elem.element_type == ElementType.TEXT_BLOCK and isinstance(elem.element, TextBlock):
            block = elem.element
            colors = [s.font.color for l in block.lines for s in l.spans if s.font.color]
            if colors:
                avg_brightness = sum(c.r + c.g + c.b for c in colors) / (3 * len(colors))
                if avg_brightness > 0.7:
                    continue
            return True

        if elem.element_type == ElementType.IMAGE:
            return True

    return False
